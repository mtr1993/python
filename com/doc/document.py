#coding=utf-8

from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage

class Document(object):
    def __init__(self, templ_docx):
        self._doc = DocxTemplate(templ_docx)
        self._pic_to_render = {}
        self._context = {}
        self._chapter = {}
        self._sequence = 0

    def sequence(self):
        self._sequence = self._sequence + 1
        return self._sequence

    def add_context(self, context = {}):
        self._context = context.copy()

    def add_heading(self, title, level=1):
        docx = self._doc.get_docx()
        docx.add_heading(title, level)

    def add_chapter(self, chapter, title = None, level=2, remove_file = True):
        try:
            docx = self._doc.get_docx()

            if title is not None:
                docx.add_heading(title, level)

            id = "chapter" + str(self.sequence())

            chapter.render()

            docx.add_paragraph("{{p " + id + "}}")

            self._chapter[id] = self._doc.new_subdoc(chapter.get_filename())

            pic_to_add = chapter.get_pic_to_add()
            if pic_to_add:
                for pic in pic_to_add:
                    path = pic_to_add[pic]["path"]
                    width = pic_to_add[pic]["width"]

                    self._pic_to_render[pic] = InlineImage(self._doc, path, width=Mm(width))
        finally:
            if remove_file:
                chapter.remove_file()

    def inlineimage(self, image_file, image_width=147):
        return InlineImage(self._doc, image_file, width=Mm(image_width))

    def render(self):
        if self._chapter or self._context:
            self._context.update(self._chapter)
            self._doc.render(self._context)

        if self._pic_to_render:
            self._doc.render(self._pic_to_render)

    def save(self, filename):
        self._doc.save(filename)
